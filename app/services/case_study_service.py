import os
import matplotlib
matplotlib.use('Agg')  # Use non-GUI backend
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import plotly.express as px
import io
import base64
from datetime import datetime
from fpdf import FPDF
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
from docx import Document
from docx.shared import Inches
import json
import re
import uuid
import requests
from app.models import db, CaseStudy, SolutionProviderInterview, ClientInterview, InviteToken
from app.services.ai_service import AIService
from app.services.metadata_service import MetadataService
from app.utils.text_processing import clean_text, detect_language
from io import BytesIO

class CaseStudyService:
    def __init__(self):
        self.output_dir = "generated_pdfs"
        os.makedirs(self.output_dir, exist_ok=True)
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        self.openai_config = {
            "model": "gpt-4",
            "temperature": 0.5,
            "top_p": 0.9,
            "presence_penalty": 0.2,
            "frequency_penalty": 0.2
        }
        self.metadata_service = MetadataService()
    
    def generate_pdf(self, case_study):
        """Generate PDF from case study and store in DB"""
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, "Case Study Report", ln=True, align='C')
            pdf.ln(10)
            pdf.set_font("Arial", "", 12)
            if case_study.final_summary:
                paragraphs = case_study.final_summary.split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        lines = paragraph.split('\n')
                        for line in lines:
                            if line.strip():
                                pdf.multi_cell(0, 5, line.strip())
                        pdf.ln(5)
            pdf_buffer = BytesIO()
            pdf.output(pdf_buffer, 'S')
            case_study.final_summary_pdf_data = pdf_buffer.getvalue()
            db.session.commit()
            return True
        except Exception as e:
            print(f"Error generating PDF: {str(e)}")
            return False
    
    def generate_word_document(self, case_study):
        """Generate Word document from case study"""
        try:
            # Create document
            doc = Document()
            
            # Set default font for the entire document
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            
            # Add case study title in bold
            title = doc.add_heading(case_study.title, 0)
            title.alignment = 1  # Center alignment
            # Set font to Arial for title
            for run in title.runs:
                run.font.name = 'Arial'
            
            # Add case study content
            if case_study.final_summary:
                paragraphs = case_study.final_summary.split('\n\n')
                
                for paragraph in paragraphs:
                    if paragraph.strip():
                        para = doc.add_paragraph(paragraph.strip())
                        # Set font to Arial for paragraph content
                        for run in para.runs:
                            run.font.name = 'Arial'
                        # Add spacing with Normal style (Arial font)
                        spacing_para = doc.add_paragraph()
                        spacing_para.style = doc.styles['Normal']
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"case_study_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # Ensure output directory exists
            os.makedirs(self.output_dir, exist_ok=True)
            
            # Save document
            doc.save(filepath)
            
            print(f"🔍 DEBUG: Word document saved to: {filepath}")
            print(f"🔍 DEBUG: File exists: {os.path.exists(filepath)}")
            
            return filepath
            
        except Exception as e:
            print(f"Error generating Word document: {str(e)}")
            return None
    
    def analyze_sentiment(self, text):
        """Analyze sentiment of text using the metadata service"""
        return self.metadata_service.analyze_sentiment(text)
    
    def generate_sentiment_chart(self, sentiment_score, output_dir="generated_pdfs"):
        """Generate sentiment visualization chart using the metadata service"""
        return self.metadata_service.generate_sentiment_chart(sentiment_score, output_dir)
    
    def extract_client_satisfaction(self, client_summary):
        """Extract client satisfaction metrics from summary using the metadata service"""
        return self.metadata_service.extract_client_satisfaction(client_summary)
    
    def generate_client_satisfaction_gauge(self, category):
        """Generate client satisfaction gauge chart"""
        try:
            # Map categories to values and colors
            category_mapping = {
                'very_satisfied': {'value': 100, 'color': '#28a745'},
                'satisfied': {'value': 75, 'color': '#17a2b8'},
                'neutral': {'value': 50, 'color': '#ffc107'},
                'dissatisfied': {'value': 25, 'color': '#fd7e14'},
                'very_dissatisfied': {'value': 0, 'color': '#dc3545'}
            }
            
            mapping = category_mapping.get(category, {'value': 50, 'color': '#ffc107'})
            
            # Create gauge chart
            fig = go.Figure(go.Indicator(
                mode = "gauge+number+delta",
                value = mapping['value'],
                domain = {'x': [0, 1], 'y': [0, 1]},
                title = {'text': "Client Satisfaction"},
                delta = {'reference': 50},
                gauge = {
                    'axis': {'range': [None, 100]},
                    'bar': {'color': mapping['color']},
                    'steps': [
                        {'range': [0, 20], 'color': "lightgray"},
                        {'range': [20, 40], 'color': "lightgray"},
                        {'range': [40, 60], 'color': "lightgray"},
                        {'range': [60, 80], 'color': "lightgray"},
                        {'range': [80, 100], 'color': "lightgray"}
                    ],
                    'threshold': {
                        'line': {'color': "red", 'width': 4},
                        'thickness': 0.75,
                        'value': 90
                    }
                }
            ))
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"satisfaction_gauge_{timestamp}.html"
            filepath = os.path.join(self.output_dir, filename)
            
            # Save chart
            fig.write_html(filepath)
            
            return filepath
            
        except Exception as e:
            print(f"Error generating satisfaction gauge: {str(e)}")
            return None
    
    def extract_client_takeaways(self, client_summary):
        """Extract key takeaways from client interview using the metadata service."""
        return self.metadata_service.extract_client_takeaways(client_summary)

    def extract_and_remove_metadata_sections(self, text, client_summary=None):
        """Extract and remove metadata sections from text using the metadata service"""
        return self.metadata_service.extract_and_remove_metadata_sections(text, client_summary)

    def create_case_study(self, user_id, title, final_summary=None):
        """Create a new case study"""
        try:
            case_study = CaseStudy(
                user_id=user_id,
                title=title,
                final_summary=final_summary
            )
            db.session.add(case_study)
            db.session.commit()
            return case_study
        except Exception as e:
            db.session.rollback()
            raise e

    def get_case_study(self, case_study_id, user_id):
        """Get a case study by ID for a specific user"""
        return CaseStudy.query.filter_by(id=case_study_id, user_id=user_id).first()

    def update_case_study(self, case_study_id, user_id, **kwargs):
        """Update a case study"""
        try:
            case_study = self.get_case_study(case_study_id, user_id)
            if not case_study:
                return None
            
            for key, value in kwargs.items():
                if hasattr(case_study, key):
                    setattr(case_study, key, value)
            
            db.session.commit()
            return case_study
        except Exception as e:
            db.session.rollback()
            raise e

    def delete_case_study(self, case_study_id, user_id):
        """Delete a case study"""
        try:
            case_study = self.get_case_study(case_study_id, user_id)
            if not case_study:
                return False
            
            db.session.delete(case_study)
            db.session.commit()
            return True
        except Exception as e:
            db.session.rollback()
            raise e

    def generate_full_case_study(self, provider_summary, client_summary, detected_language, has_client_story):
        """Generate the complete final case study with advanced features from reference"""
        try:
            if has_client_story:
                # Original prompt for when both provider and client stories exist
                full_prompt = f"""
                You are a top-tier business case study writer, creating professional, detailed, and visually attractive stories for web or PDF (inspired by Storydoc, Adobe, and top SaaS companies).
 
                IMPORTANT: Write the entire case study in {detected_language}. This includes all sections, quotes, and any additional content.
 
                Your task is to read the full Solution Provider and Client summaries below and merge them into a single, rich, multi-perspective case study. You must synthesize the insights, stories, and data into one engaging narrative.
                
                🚨 CRITICAL TITLE RULE: NEVER start your case study with "Title:" or "**Title:**" - just write the title directly on the first line!
                
                🎯 TITLE QUALITY REQUIREMENTS:
                Your title MUST be:
                - SPECIFIC to this exact case study (not generic)
                - ATTRACTIVE and compelling for business readers
                - 6-10 words maximum
                - Based on the ACTUAL solution and results described in the summaries
                - Professional yet engaging
                
                🚫 FORBIDDEN TITLE PATTERNS:
                - "Revolutionizing" anything ❌
                - "Transforming" everything ❌
                - "The Future of" anything ❌
                - Generic buzzwords like "Innovation," "Breakthrough," "Game-changer" ❌
                - Company x Client: Project Name ❌
                - Title: Your Title Here ❌
                
                ✅ EXCELLENT TITLE EXAMPLES (study these patterns):
                - "From 3 Hours to 15 Minutes: How AI Streamlined Our Onboarding"
                - "The Chatbot That Reduced Support Tickets by 70%"
                - "How We Built a System That Handles 10,000 Users Daily"
                - "The Automation That Saved Our Team 20 Hours Every Week"
                - "How Data Analytics Unlocked 40% More Revenue"
                - "The Integration That Connected 5 Systems Seamlessly"
                - "How We Scaled from 100 to 10,000 Customers in 6 Months"
                - "The Solution That Cut Processing Time from Days to Minutes"
                
                🧠 TITLE CREATION STRATEGY:
                1. Identify the SPECIFIC problem solved from the summaries
                2. Find the CONCRETE results or metrics mentioned
                3. Use ACTION words that describe what actually happened
                4. Make it SPECIFIC to this client's industry or challenge
                5. Focus on TANGIBLE outcomes, not abstract concepts
                
                💡 TITLE PATTERNS TO USE:
                - "How [Specific Solution] [Specific Result]"
                - "The [Solution Type] That [Specific Outcome]"
                - "From [Before] to [After]: How [Solution] [Result]"
                - "[Specific Metric] Improvement Through [Solution]"
                - "How [Solution] [Specific Action] [Specific Result]"
                
                CRITICAL: The first line should be ONLY the title, no prefixes, no formatting!
                Make this title so good that a business executive would want to read the case study!

                ---

                **Instructions:**
                - The **Solution Provider version is your base**; the Client version should *enhance, correct, or add* to it.
                - If the client provides a correction, update, or different number/fact for something from the provider, ALWAYS use the client's corrected version in the main story (unless it is unclear; then flag for review).
                - **IMPORTANT**: Do NOT include any "Corrected & Conflicted Replies" section in the main story. This information will be extracted separately into metadata.
                - Accuracy is CRITICAL: Double-check every fact, number, quote, and piece of information. Do NOT make any mistakes or subtle errors in the summary. Every detail must match the input summaries exactly unless you are synthesizing clearly from both. If you are unsure about a detail, do NOT invent or guess; either omit or flag it for clarification.
                - In the main story, **merge and synthesize all available details and insights** from both the Solution Provider and Client summaries: background, challenges, solutions, process, collaboration, data, quotes, and results. Do not repeat information—combine and paraphrase to build a seamless narrative.
                - **Quotes:**  
                    - Include exactly ONE impactful quote from the client in the "Customer/Client Reflection" section
                    - Include exactly ONE impactful quote from the provider in the "Testimonial/Provider Reflection" section
                    - These should be the most powerful, representative quotes
                    - Keep them concise and impactful
                - Write in clear, engaging business English. Use a mix of paragraphs, bold section headers, and bullet points.
                - Include real numbers, testimonials, collaboration stories, and unique project details whenever possible.
                - Start with a bold hero statement summarizing the main impact (NO TITLE - title will be handled separately).
                - Make each section distinct and visually scannable (use bold, bullet points, metrics, and quotes).
                - Make the results section full of specifics: show metrics, improvements, and qualitative outcomes.
                - End with a call to action for future collaboration, demo, or contact.
                - DO NOT use asterisks or Markdown stars (**) in your output. Section headers should be in ALL CAPS or plain text only.
                - DO NOT include any title or headline in the case study content - start directly with the hero statement.
                - NEVER start with "Title:" or "**Title:**" or any title formatting.
                - NEVER include company names in ALL CAPS at the beginning.
                - The content should start immediately with the story, not with any title or formatting.
                - DO NOT use section numbers like "Section 1:", "1.", "2.", etc. - just use the section headers directly.
                - Section headers should be in ALL CAPS without any numbering (e.g., "THE CHALLENGE" not "Section 1 - The Challenge").
                STRUCTURE:
 
                HERO STATEMENT / BANNER: A one-sentence summary capturing the most impactful achievement.
                INTRODUCTION
                RESEARCH AND DEVELOPMENT
                CLIENT CONTEXT AND CHALLENGES
                THE SOLUTION
                IMPLEMENTATION & COLLABORATION
                RESULTS & IMPACT
                CUSTOMER/CLIENT REFLECTION (one client quote only)
                TESTIMONIAL/PROVIDER REFLECTION (one provider quote only)
                CALL TO ACTION
 
                CONTENT RULES:
 
                - The provider's version is the base; the client's version enhances, corrects, or adds to it.
                - Use the client's corrected version if numbers or facts differ.
                - **IMPORTANT**: Do NOT include any "Corrected & Conflicted Replies" section in the main story.
                - Accuracy is critical: do not guess or invent any facts. Only use what's in the summaries.
                - Keep each section clear and scannable using ALL CAPS headers (do not bold or use markdown).
                - Main story includes exactly one quote from each side.

                Use realistic business tone and vocabulary. Do not use markdown (** **, *, -). Just clean, web/PDF-friendly output.
                
                CRITICAL: The case study content should start directly with the HERO STATEMENT. Do NOT include any title, headline, or company name formatting at the beginning of the content. If the first line looks like a title or contains company names in ALL CAPS, skip it and start with the actual story content.

                Now, here is the input:
 
                Provider Summary:
                {provider_summary}
 
                Client Summary:
                {client_summary}
                        """
            else:
                # New prompt for when only provider story exists
                full_prompt = f"""
                You are a top-tier business case study writer, creating professional, detailed, and visually attractive stories for web or PDF (inspired by Storydoc, Adobe, and top SaaS companies).
 
                IMPORTANT: Write the entire case study in {detected_language}. This includes all sections, quotes, and any additional content.
 
                Only use the Solution Provider's summary below to write a complete case study. The client did not provide input. Do not label any section with "Provider Summary" or "Title". Do not include markdown (like ** or *). Just write the case study using ALL CAPS section headers and clear business English.
                
                🚨 CRITICAL TITLE RULE: NEVER start your case study with "Title:" or "**Title:**" - just write the title directly on the first line!
                
                🎯 TITLE QUALITY REQUIREMENTS:
                Your title MUST be:
                - SPECIFIC to this exact case study (not generic)
                - ATTRACTIVE and compelling for business readers
                - 6-10 words maximum
                - Based on the ACTUAL solution and results described in the summary
                - Professional yet engaging
                
                🚫 FORBIDDEN TITLE PATTERNS:
                - "Revolutionizing" anything ❌
                - "Transforming" everything ❌
                - "The Future of" anything ❌
                - Generic buzzwords like "Innovation," "Breakthrough," "Game-changer" ❌
                - Company x Client: Project Name ❌
                - Title: Your Title Here ❌
                
                ✅ EXCELLENT TITLE EXAMPLES (study these patterns):
                - "From 3 Hours to 15 Minutes: How AI Streamlined Our Onboarding"
                - "The Chatbot That Reduced Support Tickets by 70%"
                - "How We Built a System That Handles 10,000 Users Daily"
                - "The Automation That Saved Our Team 20 Hours Every Week"
                - "How Data Analytics Unlocked 40% More Revenue"
                - "The Integration That Connected 5 Systems Seamlessly"
                - "How We Scaled from 100 to 10,000 Customers in 6 Months"
                - "The Solution That Cut Processing Time from Days to Minutes"
                
                🧠 TITLE CREATION STRATEGY:
                1. Identify the SPECIFIC problem solved from the summary
                2. Find the CONCRETE results or metrics mentioned
                3. Use ACTION words that describe what actually happened
                4. Make it SPECIFIC to this client's industry or challenge
                5. Focus on TANGIBLE outcomes, not abstract concepts
                
                💡 TITLE PATTERNS TO USE:
                - "How [Specific Solution] [Specific Result]"
                - "The [Solution Type] That [Specific Outcome]"
                - "From [Before] to [After]: How [Solution] [Result]"
                - "[Specific Metric] Improvement Through [Solution]"
                - "How [Solution] [Specific Action] [Specific Result]"
                
                CRITICAL: The first line should be ONLY the title, no prefixes, no formatting!
                Make this title so good that a business executive would want to read the case study!
 
                STRUCTURE:
 
                HERO STATEMENT / BANNER: A one-sentence summary capturing the most impactful achievement.
                INTRODUCTION
                RESEARCH AND DEVELOPMENT
                CLIENT CONTEXT AND CHALLENGES
                THE SOLUTION
                IMPLEMENTATION & COLLABORATION
                RESULTS & IMPACT
                TESTIMONIAL/PROVIDER REFLECTION (one quote from the provider)
                CALL TO ACTION
 
                CONTENT RULES:
 
                - Maintain credibility: do not fabricate specific client claims, only rephrase insights from the provider.
                - Keep each section clear and scannable using ALL CAPS headers (no bolding or markdown).
                - Include one quote in each reflection section.
                - DO NOT include any title or headline in the case study content - start directly with the hero statement.
                - NEVER start with "Title:" or "**Title:**" or any title formatting.
                - NEVER include company names in ALL CAPS at the beginning.
                - The content should start immediately with the story, not with any title or formatting.
                - DO NOT use section numbers like "Section 1:", "1.", "2.", etc. - just use the section headers directly.
                - Section headers should be in ALL CAPS without any numbering (e.g., "THE CHALLENGE" not "Section 1 - The Challenge").

                Use a realistic tone and avoid generic phrases. Just output the full case study without section labels, markdown, or references to instructions.
                
                CRITICAL: The case study content should start directly with the HERO STATEMENT. Do NOT include any title, headline, or company name formatting at the beginning of the content. If the first line looks like a title or contains company names in ALL CAPS, skip it and start with the actual story content.
 
                Now, here is the input:

                Provider Summary:
                {provider_summary}
                """

            headers = {
                "Authorization": f"Bearer {self.openai_api_key}",
                "Content-Type": "application/json"
            }

            payload = {
                "model": self.openai_config["model"],
                "messages": [
                    {"role": "system", "content": full_prompt},
                ],
                "temperature": 0.5,
                "top_p": 0.9
            }

            response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload,timeout=120)
            result = response.json()
            case_study_text = result["choices"][0]["message"]["content"]
            cleaned = clean_text(case_study_text)

            main_story, meta_data = self.extract_and_remove_metadata_sections(cleaned, client_summary)
            
            # Remove any title-like content from the beginning of the case study
            lines = main_story.split('\n')
            if lines:
                first_line = lines[0].strip()
                print(f"🔍 DEBUG: First line of case study content: '{first_line}'")
                # If first line looks like a title (contains ALL CAPS words, company names, or "Title:"), remove it
                if (first_line.isupper() or 
                    ':' in first_line or 
                    first_line.startswith('**Title:**') or
                    first_line.startswith('Title:') or
                    any(word.isupper() and len(word) > 3 for word in first_line.split()) or
                    'TRANSFORMS' in first_line or 'REVOLUTIONIZING' in first_line):
                    print(f"🔍 WARNING: Removing title-like first line: '{first_line}'")
                    lines = lines[1:]
                    main_story = '\n'.join(lines).strip()
                    print(f"🔍 FIXED: Case study now starts with: '{lines[0].strip() if lines else 'None'}'")
                # Also check for "Title:" format and remove it
                elif first_line.startswith('Title:'):
                    print(f"🔍 WARNING: Removing 'Title:' prefix from: '{first_line}'")
                    # Extract the actual title content after "Title:"
                    actual_title = first_line.replace('Title:', '').strip()
                    lines[0] = actual_title
                    main_story = '\n'.join(lines).strip()
                    print(f"🔍 FIXED: Removed 'Title:' prefix, now: '{actual_title}'")
                # Check for any line that contains "Title:" anywhere in it
                elif 'Title:' in first_line:
                    print(f"🔍 WARNING: Found 'Title:' in line: '{first_line}'")
                    # Extract the actual title content after "Title:"
                    actual_title = first_line.split('Title:', 1)[1].strip()
                    lines[0] = actual_title
                    main_story = '\n'.join(lines).strip()
                    print(f"🔍 FIXED: Removed 'Title:' prefix, now: '{actual_title}'")
                else:
                    print(f"🔍 OK: First line looks good: '{first_line}'")
            
            # Remove any remaining quotes and asterisks from the entire content
            main_story = main_story.replace('"', '').replace('**', '').replace('*', '')
            print(f"🔍 CLEANED: Removed quotes and asterisks from case study content")
            
            # Generate corrected & conflicted replies metadata if client story exists
            if has_client_story and client_summary:
                corrected_replies = self.generate_corrected_conflicted_replies(provider_summary, client_summary)
                meta_data["corrected_conflicted_replies"] = corrected_replies
            
            print("Meta data being saved:", meta_data)
            return main_story, meta_data

        except Exception as e:
            print(f"Error generating full case study: {str(e)}")
            raise e

    def draft_quote_from_summary(self, summary, speaker="Client"):
        """Simple template-based fallback if OpenAI is not available"""
        return self.metadata_service.draft_quote_from_summary(summary, speaker)

    def analyze_client_sentiment(self, client_summary):
        """Analyze sentiment of client summary using the metadata service"""
        return self.metadata_service.analyze_sentiment(client_summary)
    
    def generate_corrected_conflicted_replies(self, provider_summary, client_summary):
        """Generate corrected and conflicted replies metadata"""
        try:
            if not self.openai_api_key:
                return "AI service not available for generating corrected replies."
            
            prompt = f"""
            Analyze the provider and client summaries below and identify any corrections, contradictions, or additional information provided by the client.
            
            Focus on:
            - Numbers, metrics, or facts that differ between provider and client
            - Information that the client added that wasn't mentioned by the provider
            - Corrections to the provider's version
            - Additional context or details provided by the client
            
            Provider Summary:
            {provider_summary}
            
            Client Summary:
            {client_summary}
            
            Generate a concise list of corrected and conflicted replies. Format as bullet points.
            If there are no corrections or conflicts, return "No corrections or conflicts identified."
            """
            
            headers = {
                "Authorization": f"Bearer {self.openai_api_key}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": self.openai_config["model"],
                "messages": [{"role": "system", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": 500
            }
            
            response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload,timeout=120)
            result = response.json()
            
            if "choices" in result and len(result["choices"]) > 0:
                return result["choices"][0]["message"]["content"].strip()
            else:
                return "Failed to generate corrected replies."
                
        except Exception as e:
            print(f"Error generating corrected replies: {str(e)}")
            return "Error generating corrected replies." 
