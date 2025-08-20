from flask import Blueprint, render_template, request, jsonify, session, redirect, url_for, send_from_directory, send_file
import uuid
import os
import requests
import json
from datetime import datetime, timedelta, UTC
from app.models import db, CaseStudy, SolutionProviderInterview, ClientInterview, InviteToken
from app.utils.text_processing import clean_text, detect_language
from app.utils.auth_helpers import get_current_user_id
from app.services.ai_service import AIService
from flasgger import swag_from
import logging

logger = logging.getLogger(__name__)

bp = Blueprint('main', __name__)

# API Configuration
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

openai_config = {
    "model": "gpt-4",
    "temperature": 0.5,
    "top_p": 0.9,
    "presence_penalty": 0.2,
    "frequency_penalty": 0.2
}

@bp.route("/")
def serve_index():
    """Serve the main index page - redirect to dashboard if logged in, otherwise to login"""
    user_id = get_current_user_id()
    if user_id:
        return redirect(url_for('main.dashboard'))
    else:
        return redirect(url_for('main.login'))

@bp.route("/favicon.ico")
def favicon():
    """Serve favicon to prevent 404 errors"""
    return "", 204  # Return no content to prevent 404 errors

@bp.route("/login")
def login():
    """Serve the login page"""
    return render_template('login.html')

@bp.route("/signup")
def signup():
    """Serve the signup page"""
    return render_template('signup.html')

@bp.route("/verification")
def verification(): 
    """Serve the verification page"""    
    return render_template('verification.html')

@bp.route("/dashboard")
def dashboard():
    """Serve the dashboard page - requires authentication"""
    user_id = get_current_user_id()
    if not user_id:
        return redirect(url_for('main.login'))
    return render_template('dashboard.html')

@bp.route("/feedback")
def feedback():
    """Serve the feedback page - requires authentication"""
    user_id = get_current_user_id()
    if not user_id:
        return redirect(url_for('main.login'))
    return render_template('feedback.html')

@bp.route("/client")
def client():
    """Serve the client page - requires authentication"""
    user_id = get_current_user_id()
    if not user_id:
        return redirect(url_for('main.login'))
    return render_template('client.html')

@bp.route("/index")
def index_page():
    return render_template('index.html')

@bp.route("/session")
@swag_from({
    'tags': ['Real-time Interviews'],
    'summary': 'Create OpenAI realtime session',
    'description': 'Create OpenAI realtime session for voice interviews',
    'responses': {
        200: {
            'description': 'Session configuration for real-time voice chat',
            'content': {
                'application/json': {
                    'schema': {
                        'type': 'object',
                        'properties': {
                            'client_secret': {
                                'type': 'object',
                                'properties': {
                                    'value': {'type': 'string'}
                                }
                            }
                        }
                    }
                }
            }
        }
    }
})
def create_session():
    """Create OpenAI realtime session"""
    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json",
    }
    data = {
        "model": "gpt-4o-realtime-preview-2024-12-17",
        "voice": "coral"
    }
    response = requests.post("https://api.openai.com/v1/realtime/sessions", headers=headers, json=data)
    return jsonify(response.json())

@bp.route("/save_transcript", methods=["POST"])
@swag_from({
    'tags': ['Real-time Interviews'],
    'summary': 'Save provider interview transcript',
    'description': 'Save solution provider interview transcript',
    'parameters': [
        {
            'name': 'provider_session_id',
            'in': 'query',
            'required': True,
            'schema': {'type': 'string'}
        }
    ],
    'requestBody': {
        'required': True,
        'content': {
            'application/json': {
                'schema': {
                    'type': 'array',
                    'items': {
                        'type': 'object',
                        'properties': {
                            'speaker': {'type': 'string'},
                            'text': {'type': 'string'}
                        }
                    }
                }
            }
        }
    },
    'responses': {
        200: {
            'description': 'Transcript saved to DB',
            'content': {
                'application/json': {
                    'schema': {
                        'type': 'object',
                        'properties': {
                            'status': {'type': 'string'},
                            'message': {'type': 'string'}
                        }
                    }
                }
            }
        },
        400: {'description': 'Missing provider_session_id'}
    }
})
def save_transcript():
    """Save provider transcript"""
    try:
        raw_transcript = request.get_json()
        transcript_lines = []
        buffer = {"ai": "", "user": ""}
        last_speaker = None

        for entry in raw_transcript:
            speaker = entry.get("speaker", "").lower()
            text = entry.get("text", "").strip()
            if not text:
                continue
            if speaker != last_speaker and last_speaker is not None:
                if buffer[last_speaker]:
                    transcript_lines.append(f"{last_speaker.upper()}: {buffer[last_speaker].strip()}")
                    buffer[last_speaker] = ""
            buffer[speaker] += " " + text
            last_speaker = speaker

        if last_speaker and buffer[last_speaker]:
            transcript_lines.append(f"{last_speaker.upper()}: {buffer[last_speaker].strip()}")

        full_transcript = "\n".join(transcript_lines)

        # Get provider_session_id from query string
        provider_session_id = request.args.get("provider_session_id")
        if not provider_session_id:
            return jsonify({"status": "error", "message": "Missing provider_session_id"}), 400

        # Store in DB
        interview = SolutionProviderInterview.query.filter_by(session_id=provider_session_id).first()
        if interview:
            interview.transcript = full_transcript
            db.session.commit()

        return jsonify({"status": "success", "message": "Transcript saved to DB"})

    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500

@bp.route("/save_client_transcript", methods=["POST"])
@swag_from({
    'tags': ['Real-time Interviews'],
    'summary': 'Save client interview transcript',
    'description': 'Save client interview transcript',
    'parameters': [
        {
            'name': 'token',
            'in': 'query',
            'required': True,
            'schema': {'type': 'string'}
        }
    ],
    'requestBody': {
        'required': True,
        'content': {
            'application/json': {
                'schema': {
                    'type': 'array',
                    'items': {
                        'type': 'object',
                        'properties': {
                            'speaker': {'type': 'string'},
                            'text': {'type': 'string'}
                        }
                    }
                }
            }
        }
    },
    'responses': {
        200: {
            'description': 'Client transcript saved',
            'content': {
                'application/json': {
                    'schema': {
                        'type': 'object',
                        'properties': {
                            'status': {'type': 'string'},
                            'message': {'type': 'string'},
                            'session_id': {'type': 'string'}
                        }
                    }
                }
            }
        },
        400: {'description': 'Missing token'},
        404: {'description': 'Invalid token'}
    }
})
def save_client_transcript():
    """Save client transcript"""
    try:
        raw_transcript = request.get_json()
        transcript_lines = []
        buffer = {"ai": "", "user": ""}
        last_speaker = None

        for entry in raw_transcript:
            speaker = entry.get("speaker", "").lower()
            text = entry.get("text", "").strip()
            if not text:
                continue
            if speaker != last_speaker and last_speaker is not None:
                if buffer[last_speaker]:
                    transcript_lines.append(f"{last_speaker.upper()}: {buffer[last_speaker].strip()}")
                    buffer[last_speaker] = ""
            buffer[speaker] += " " + text
            last_speaker = speaker

        if last_speaker and buffer[last_speaker]:
            transcript_lines.append(f"{last_speaker.upper()}: {buffer[last_speaker].strip()}")

        full_transcript = "\n".join(transcript_lines)

        # Get token from query string
        token = request.args.get("token")
        if not token:
            return jsonify({"status": "error", "message": "Missing token"}), 400

        # Get case_study_id from token
        invite = InviteToken.query.filter_by(token=token).first()
        if not invite:
            return jsonify({"status": "error", "message": "Invalid token"}), 404

        # Create or update ClientInterview
        client_session_id = str(uuid.uuid4())
        interview = ClientInterview.query.filter_by(case_study_id=invite.case_study_id).first()

        if interview:
            interview.transcript = full_transcript
        else:
            interview = ClientInterview(
                case_study_id=invite.case_study_id,
                session_id=client_session_id,
                transcript=full_transcript
            )
            db.session.add(interview)

        db.session.commit()
        return jsonify({"status": "success", "message": "Client transcript saved", "session_id": client_session_id})

    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500

@bp.route("/generate_summary", methods=["POST"])
@swag_from({
    'tags': ['Interviews'],
    'summary': 'Generate interview summary',
    'description': 'Generate a structured case study summary from interview transcript',
    'requestBody': {
        'required': True,
        'content': {
            'application/json': {
                'schema': {
                    'type': 'object',
                    'required': ['transcript'],
                    'properties': {
                        'transcript': {
                            'type': 'string',
                            'description': 'Interview transcript text'
                        }
                    }
                }
            }
        }
    },
    'responses': {
        200: {
            'description': 'Summary generated successfully',
            'content': {
                'application/json': {
                    'schema': {
                        'type': 'object',
                        'properties': {
                            'status': {'type': 'string'},
                            'summary': {'type': 'string'}
                        }
                    }
                }
            }
        },
        400: {'description': 'Transcript is missing'},
        500: {'description': 'Internal server error'}
    }
})
def generate_summary():
    """Generate summary from transcript"""
    print("🚨🚨🚨 NEW CODE IS RUNNING - POST-PROCESSING SHOULD WORK! 🚨🚨🚨")
    try:
        data = request.get_json()
        transcript = data.get("transcript", "")

        if not transcript:
            return jsonify({"status": "error", "message": "Transcript is missing."}), 400

        # Detect language from transcript
        detected_language = detect_language(transcript)
        
        # Use the detected language in the prompt
        prompt = f"""
        You are a professional case study writer. Your job is to generate a **rich, structured, human-style business case study** from a transcript of a real voice interview.

        IMPORTANT: Write the entire case study in {detected_language}. This includes all sections, quotes, and any additional content.
        This is an external project: the speaker is the solution provider describing a project they delivered to a client. Your task is to write a clear, emotionally intelligent case study from their perspective—based ONLY on what's in the transcript.

        🚨 CRITICAL TITLE RULE: NEVER start your case study with "Title:" or "**Title:**" - just write the title directly on the first line!

        --- 

        ❌ DO NOT INVENT ANYTHING  
        - Do NOT fabricate dialogue or add made-up details  
        - Do NOT simulate the interview format  
        - Do NOT assume or imagine info not explicitly said  

        ✅ USE ONLY what's really in the transcript. If a piece of information (like a client quote) wasn't provided, craft a brief, realistic-sounding quote that captures the client's sentiment based on what they did say.

        --- 

        ### ✍️ CASE STUDY STRUCTURE (MANDATORY)

        Title (first line only—no extra formatting): 
        CRITICAL: Your title MUST be a short hook like these examples:
        - How AI Transformed Customer Service
        - The Chatbot That Changed Everything
        - Automation That Saved the Day
        - The Solution That Locked Success
        
        FORBIDDEN FORMATS:
        - Company x Client: Project Name ❌
        - Provider x Partner: Solution ❌  
        - TechYard x Vostro: Lock it up ❌
        - Title: Your Title Here ❌
        - **Title:** Your Title Here ❌
        - Title: Your Title Here ❌
        
        Your title must be 5-8 words, no company names, just a catchy hook!
        DO NOT write "Title:" before your title - just write the title directly!
        NEVER start with "Title:" or "**Title:**" - just write the title!
        CRITICAL: The first line should be ONLY the title, no prefixes, no formatting!

        --- 

        Hero Paragraph (no header)  
        3–4 sentences introducing the client, their industry, and their challenge; then introduce the provider and summarize the delivery.

        --- 

        THE CHALLENGE  
        - What problem was the client solving?  
        - Why was it important?  
        - Any context on scale, goals, or mission

        --- 

        THE SOLUTION  
        - Describe the delivered product/service/strategy  
        - Break down key components and clever features

        --- 

        IMPLEMENTATION & COLLABORATION  
        - How was it rolled out?  
        - What was the teamwork like?  
        - Any turning points or lessons learned

        --- 

        RESULTS & IMPACT  
        - What changed for the client?  
        - Include any real metrics (e.g., 40% faster onboarding)  
        - Mention qualitative feedback if shared

        --- 

        CLIENT QUOTE  
        - If the transcript contains a direct, verbatim quote from the client or solution provider, include it as spoken.  
        - If no direct quote is present, compose one elegant sentence from the client's or provider's perspective. Use only language, tone, and key points found in the transcript to craft a testimonial that feels genuine, highlights the solution's impact, and reads like a professional endorsement.

        --- 

        REFLECTIONS & CLOSING  
        - What did this mean for the provider's team?  
        - End with a warm, forward-looking sentence.

        --- 

        ❌ Do NOT use section numbers like "Section 1:", "1.", "2.", etc. - just use the section headers directly.
        ❌ Section headers should be in ALL CAPS without any numbering (e.g., "THE CHALLENGE" not "Section 1 - The Challenge").
        
        🎯 GOAL:  
        A vivid, accurate, human-sounding case study grounded entirely in the transcript.

        Transcript:
        {transcript}
        """

        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": openai_config["model"],
            "messages": [{"role": "system", "content": prompt}],
            "temperature": openai_config["temperature"],
            "top_p": openai_config["top_p"],
            "presence_penalty": openai_config["presence_penalty"],
            "frequency_penalty": openai_config["frequency_penalty"]
        }

        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
        result = response.json()
        case_study = result["choices"][0]["message"]["content"]
        cleaned = clean_text(case_study)
        
        print(f"🔍 TEST: Application is using latest code - post-processing should work now!")
        print(f"🔍 TEST: Raw AI response first line: '{case_study.split('\n')[0] if case_study else 'None'}'")
        print(f"🚨🚨🚨 POST-PROCESSING LOGIC SHOULD EXECUTE NOW! 🚨🚨🚨")
        
        # Post-process to ensure title is not in old format
        lines = cleaned.split('\n')
        if lines:
            first_line = lines[0].strip()
            print(f"🔍 DEBUG: AI generated first line: '{first_line}'")
            
            # If first line looks like old format, replace it with a generic hook
            if (' x ' in first_line and ':' in first_line) or ('X' in first_line and ':' in first_line):
                print(f"🔍 WARNING: AI generated old format title: '{first_line}'")
                # Replace with a generic hook based on the content
                if 'chatbot' in first_line.lower() or 'ai' in first_line.lower():
                    lines[0] = "How AI Transformed Customer Service"
                elif 'automation' in first_line.lower():
                    lines[0] = "Automation That Changed Everything"
                elif 'lock' in first_line.lower():
                    lines[0] = "The Solution That Locked Success"
                else:
                    lines[0] = "A Success Story That Transformed Business"
                cleaned = '\n'.join(lines)
                print(f"🔍 FIXED: Replaced with: '{lines[0]}'")
            # Also check for "Title:" format and remove it
            elif first_line.startswith('Title:'):
                print(f"🔍 WARNING: Removing 'Title:' prefix from: '{first_line}'")
                # Extract the actual title content after "Title:"
                actual_title = first_line.replace('Title:', '').strip()
                lines[0] = actual_title
                cleaned = '\n'.join(lines)
                print(f"🔍 FIXED: Removed 'Title:' prefix, now: '{actual_title}'")
            # Check for any line that contains "Title:" anywhere in it
            elif 'Title:' in first_line:
                print(f"🔍 WARNING: Found 'Title:' in line: '{first_line}'")
                # Extract the actual title content after "Title:"
                actual_title = first_line.split('Title:', 1)[1].strip()
                lines[0] = actual_title
                cleaned = '\n'.join(lines)
                print(f"🔍 FIXED: Removed 'Title:' prefix, now: '{actual_title}'")
            else:
                print(f"🔍 OK: Title looks good: '{first_line}'")
        
        # Remove any quotes and asterisks from the entire content
        cleaned = cleaned.replace('"', '').replace('**', '').replace('*', '')
        print(f"🔍 CLEANED: Removed quotes and asterisks from case study content")
        
        # Extract names using AI service
        ai_service = AIService()
        names = ai_service.extract_names_from_case_study(cleaned)
        
        # First save to DB and get case_study_id
        provider_session_id = str(uuid.uuid4())  # Generate a session ID now
        case_study_id = store_solution_provider_session(provider_session_id, cleaned)

        return jsonify({
            "status": "success",
            "text": cleaned,
            "names": names,
            "provider_session_id": provider_session_id,
            "case_study_id": case_study_id
        })

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@bp.route("/save_provider_summary", methods=["POST"])
def save_provider_summary():
    """Save provider summary"""
    try:
        data = request.get_json()
        provider_session_id = data.get("provider_session_id")
        updated_summary = data.get("summary")

        if not provider_session_id or not updated_summary:
            return jsonify({"status": "error", "message": "Missing data"}), 400

        # Get interview from DB
        interview = SolutionProviderInterview.query.filter_by(session_id=provider_session_id).first()
        if not interview:
            return jsonify({"status": "error", "message": "Session not found"}), 404

        # Update summary
        interview.summary = updated_summary

        # Extract names from the new summary (but don't change the title format)
        ai_service = AIService()
        names = ai_service.extract_names_from_case_study(updated_summary)
        
        # Keep the existing title as is - it should be a short hook, not the old format
        # The title was already generated correctly in the initial summary generation
        case_study = CaseStudy.query.filter_by(id=interview.case_study_id).first()
        if case_study:
            # Don't override the title - keep the short hook that was already generated
            pass

        db.session.commit()

        return jsonify({
            "status": "success",
            "message": "Summary and title updated",
            "names": names,
            "case_study_id": case_study.id,
            "provider_session_id": provider_session_id
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500

@bp.route("/generate_client_summary", methods=["POST"])
@swag_from({
    'tags': ['Interviews'],
    'summary': 'Generate client interview summary',
    'description': 'Generate a client perspective summary from client interview transcript',
    'parameters': [
        {
            'name': 'token',
            'in': 'query',
            'required': True,
            'schema': {'type': 'string'},
            'description': 'Client interview token'
        }
    ],
    'requestBody': {
        'required': True,
        'content': {
            'application/json': {
                'schema': {
                    'type': 'object',
                    'required': ['transcript'],
                    'properties': {
                        'transcript': {
                            'type': 'string',
                            'description': 'Client interview transcript text'
                        }
                    }
                }
            }
        }
    },
    'responses': {
        200: {
            'description': 'Client summary generated successfully',
            'content': {
                'application/json': {
                    'schema': {
                        'type': 'object',
                        'properties': {
                            'status': {'type': 'string'},
                            'summary': {'type': 'string'}
                        }
                    }
                }
            }
        },
        400: {'description': 'Transcript or token missing'},
        500: {'description': 'Internal server error'}
    }
})
def generate_client_summary():
    """Generate client summary from transcript"""
    try:
        data = request.get_json()
        transcript = data.get("transcript", "")
        token = request.args.get("token")

        if not transcript:
            return jsonify({"status": "error", "message": "Transcript is missing."}), 400
        if not token:
            return jsonify({"status": "error", "message": "Missing token"}), 400
            
        detected_language = detect_language(transcript)

        prompt = f"""
You are a professional case study writer. Your job is to generate a **rich, human-style client perspective** on a project delivered by a solution provider.
IMPORTANT: Write the entire case study in {detected_language}. This includes all sections, quotes, and any additional content.
        - DO NOT include the transcript itself in the output.

This is a **client voice** case study — the transcript you're given is from the client who received the solution. You will create a short, structured reflection based entirely on what they shared.

---

✅ Use only the information provided in the transcript  
❌ Do NOT invent or assume missing details
❌ Do NOT use section numbers like "Section 1:", "1.", "2.", etc. - just use the section headers directly

---

### Structure:

**PROJECT REFLECTION (Client Voice)**  
A warm, professional 3–5 sentence paragraph that shares:  
- What the project was  
- What the client's experience was like  
- The results or value they got  
- A light personal note if they gave one

---

**CLIENT QUOTE**  
Include a short quote from the client (verbatim if given, otherwise craft one from the content).  
Make it feel authentic, appreciative, and aligned with their actual words.

---

🎯 GOAL:  
Provide a simple, balanced, human-sounding reflection from the client that complements the full case study.

Transcript:
{transcript}
"""

        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": openai_config["model"],
            "messages": [
                {"role": "system", "content": prompt},
            ],
            "temperature": openai_config["temperature"],
            "top_p": openai_config["top_p"],
            "presence_penalty": openai_config["presence_penalty"],
            "frequency_penalty": openai_config["frequency_penalty"]
        }

        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
        result = response.json()
        summary = result["choices"][0]["message"]["content"]
        cleaned = clean_text(summary)

        invite = InviteToken.query.filter_by(token=token).first()
        if not invite:
            return jsonify({"status": "error", "message": "Invalid token"}), 404

        client_interview = ClientInterview.query.filter_by(case_study_id=invite.case_study_id).first()
        if client_interview:
            client_interview.summary = cleaned
            db.session.commit()

        return jsonify({
            "status": "success",
            "text": cleaned,
            "case_study_id": invite.case_study_id
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500

@bp.route("/client-interview/<token>", methods=["GET"])
@swag_from({
    'tags': ['Interviews'],
    'summary': 'Get client interview data',
    'description': 'Get client interview data using token',
    'parameters': [
        {
            'name': 'token',
            'in': 'path',
            'required': True,
            'schema': {'type': 'string'},
            'description': 'Client interview token'
        }
    ],
    'responses': {
        200: {
            'description': 'Client interview data retrieved successfully',
            'content': {
                'application/json': {
                    'schema': {
                        'type': 'object',
                        'properties': {
                            'status': {'type': 'string'},
                            'provider_name': {'type': 'string'},
                            'client_name': {'type': 'string'},
                            'project_name': {'type': 'string'},
                            'provider_summary': {'type': 'string'}
                        }
                    }
                }
            }
        },
        404: {'description': 'Invalid or expired token, or case study not found'}
    }
})
def client_interview(token):
    """Get client interview data"""
    try:
        # Fetch InviteToken by token
        invite = InviteToken.query.filter_by(token=token).first()
        if not invite or invite.used:
            return jsonify({"status": "error", "message": "Invalid or expired link"}), 404

        # Fetch CaseStudy and linked SolutionProviderInterview
        case_study = CaseStudy.query.filter_by(id=invite.case_study_id).first()
        if not case_study:
            return jsonify({"status": "error", "message": "Case study not found"}), 404

        provider_interview = case_study.solution_provider_interview
        if not provider_interview:
            return jsonify({"status": "error", "message": "Provider interview not found"}), 404

        # Mark the invite token as used
        invite.used = True
        db.session.commit()

        # Extract names from FINAL SUMMARY (not provider interview) to get updated names
        ai_service = AIService()
        if case_study.final_summary:
            # Use final summary for name extraction to get the updated/edited names
            extracted_names = ai_service.extract_names_from_case_study(case_study.final_summary)
            provider_name = extracted_names.get("lead_entity", "Unknown")
            client_name = extracted_names.get("partner_entity", "")
            project_name = extracted_names.get("project_title", "Unknown Project")
            provider_summary = case_study.final_summary  # Use final summary content
            print(f"✅ Client interview using FINAL SUMMARY names: {extracted_names}")
        else:
            # Fallback to provider interview if no final summary exists
            extracted_names = ai_service.extract_names_from_case_study(provider_interview.summary)
            provider_name = extracted_names.get("lead_entity", "Unknown")
            client_name = extracted_names.get("partner_entity", "")
            project_name = extracted_names.get("project_title", "Unknown Project")
            provider_summary = provider_interview.summary
            print(f"⚠️ Client interview using PROVIDER INTERVIEW names (no final summary): {extracted_names}")

        return jsonify({
            "status": "success",
            "provider_name": provider_name,
            "client_name": client_name,
            "project_name": project_name,
            "provider_summary": provider_summary
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500

@bp.route("/client/<token>")
def serve_client_interview(token):
    """Serve client interview page"""
    return render_template('client.html')

@bp.route("/generate_client_interview_link", methods=["POST"])
@swag_from({
    'tags': ['Interviews'],
    'summary': 'Generate client interview link',
    'description': 'Generate a unique link for client interview participation',
    'requestBody': {
        'required': True,
        'content': {
            'application/json': {
                'schema': {
                    'type': 'object',
                    'required': ['case_study_id'],
                    'properties': {
                        'case_study_id': {
                            'type': 'integer',
                            'description': 'ID of the case study'
                        },
                        'solution_provider': {
                            'type': 'string',
                            'description': 'Name of the solution provider'
                        },
                        'client_name': {
                            'type': 'string',
                            'description': 'Name of the client'
                        },
                        'project_name': {
                            'type': 'string',
                            'description': 'Name of the project'
                        }
                    }
                }
            }
        }
    },
    'responses': {
        200: {
            'description': 'Client interview link generated successfully',
            'content': {
                'application/json': {
                    'schema': {
                        'type': 'object',
                        'properties': {
                            'status': {'type': 'string'},
                            'link': {'type': 'string'},
                            'token': {'type': 'string'}
                        }
                    }
                }
            }
        },
        400: {'description': 'Case study ID is required'},
        404: {'description': 'Case study not found'}
    }
})
def generate_client_interview_link():
    """Generate client interview link"""
    try:
        data = request.get_json()
        case_study_id = data.get("case_study_id")
        # Get edited names from frontend if provided
        solution_provider = data.get("solution_provider")
        client_name = data.get("client_name") 
        project_name = data.get("project_name")
        
        if not case_study_id:
            return jsonify({"status": "error", "message": "Missing case_study_id."}), 400

        # Make sure this case study exists
        case_study = CaseStudy.query.filter_by(id=case_study_id).first()
        if not case_study:
            return jsonify({"status": "error", "message": "Invalid case study ID."}), 400

        # ONLY use final summary - no fallback to provider summary
        if not case_study.final_summary:
            return jsonify({"status": "error", "message": "Final summary is required to generate client interview link. Please complete the final story first."}), 400

        # Use edited names from frontend if provided, otherwise extract from final summary
        if solution_provider and client_name and project_name:
            extracted_names = {
                "lead_entity": solution_provider,
                "partner_entity": client_name,
                "project_title": project_name
            }
            print(f"✅ Using edited names from frontend: {extracted_names}")
        else:
            # Extract names from final summary only
            ai_service = AIService()
            extracted_names = ai_service.extract_names_from_case_study(case_study.final_summary)
            print(f"✅ Using extracted names from final summary: {extracted_names}")

        token = create_client_session(case_study_id)
        if not token:
            return jsonify({"status": "error", "message": "Failed to create client session."}), 500
        
        BASE_URL = os.getenv("BASE_URL", "http://127.0.0.1:10000")
        interview_link = f"{BASE_URL}/client/{token}"
        
        print(f"🔗 Generated interview link: {interview_link}")
        print(f"🔗 BASE_URL: {BASE_URL}")
        print(f"🔗 Token: {token}")
        
        # Save to database
        provider_interview = SolutionProviderInterview.query.filter_by(case_study_id=case_study_id).first()
        if provider_interview:
            provider_interview.client_link_url = interview_link
            db.session.commit()
            print(f"✅ Saved client_link_url to database: {interview_link}")
        else:
            print(f"❌ No provider interview found for case_study_id: {case_study_id}")

        return jsonify({"status": "success", "interview_link": interview_link})
    except Exception as e:
        db.session.rollback()
        print(f"❌ Error in generate_client_interview_link: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 500

@bp.route("/get_provider_transcript", methods=["GET"])
@swag_from({
    'tags': ['Interviews'],
    'summary': 'Get provider transcript',
    'description': 'Retrieve the solution provider interview transcript',
    'parameters': [
        {
            'name': 'token',
            'in': 'query',
            'required': True,
            'schema': {'type': 'string'},
            'description': 'Client interview token'
        }
    ],
    'responses': {
        200: {
            'description': 'Provider transcript retrieved successfully',
            'content': {
                'application/json': {
                    'schema': {
                        'type': 'object',
                        'properties': {
                            'status': {'type': 'string'},
                            'transcript': {'type': 'string'}
                        }
                    }
                }
            }
        },
        400: {'description': 'Missing token'},
        404: {'description': 'Invalid token or transcript not found'}
    }
})
def get_provider_transcript():
    """Get provider transcript"""
    try:
        token = request.args.get("token")
        if not token:
            return jsonify({"status": "error", "message": "Missing token"}), 400

        # Get case_study_id from token
        invite = InviteToken.query.filter_by(token=token).first()
        if not invite:
            return jsonify({"status": "error", "message": "Invalid token"}), 404

        # Get the provider interview transcript
        provider_interview = SolutionProviderInterview.query.filter_by(case_study_id=invite.case_study_id).first()
        if not provider_interview or not provider_interview.transcript:
            return jsonify({"status": "error", "message": "Provider transcript not found"}), 404

        return jsonify({
            "status": "success",
            "transcript": provider_interview.transcript
        })

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500



# @bp.route("/extract_names", methods=["POST"]) - REMOVED: Duplicate endpoint, using /api/extract_names instead
# @swag_from({
#     'tags': ['Case Studies'],
#     'summary': 'Extract names from case study',
#     'description': 'Extract company and project names from case study text',
#     'requestBody': {
#         'required': True,
#         'content': {
#             'application/json': {
#                 'schema': {
#                     'type': 'object',
#                     'required': ['summary'],
#                     'properties': {
#                         'summary': {'type': 'string'}
#                     }
#                 }
#             }
#         }
#     },
#     'responses': {
#         200: {
#             'description': 'Names extracted successfully',
#             'content': {
#                 'application/json': {
#                     'schema': {
#                         'type': 'object',
#                         'properties': {
#                             'status': {'type': 'string'},
#                             'names': {'type': 'string'},
#                             'method': {'type': 'string'}
#                         }
#                     }
#                 }
#             }
#         },
#         400: {'description': 'Bad Request'}
#     }
# })
# def extract_names():
#     """Extract names from case study text"""
#     try:
#         data = request.get_json()
#         summary = data.get("summary", "")
#         if not summary:
#             return jsonify({"status": "error", "message": "Missing summary"}), 400
# 
#         print(f"🎯 Starting name extraction for summary length: {len(summary)}")
#         ai_service = AIService()
#         names = ai_service.extract_names_from_case_study(summary)
#         print(f"🎯 Name extraction result: {names}")
#         
#         return jsonify({
#             "status": "success", 
#             "names": names,
#             "method": "llm"  # Add method indicator
#         })
#     except Exception as e:
#         print(f"❌ Error in extract_names endpoint: {str(e)}")
#         return jsonify({"status": "error", "message": str(e)}), 500

# @bp.route("/save_final_summary", methods=["POST"]) - REMOVED: Duplicate endpoint, using /api/save_final_summary instead
# @swag_from({
#     'tags': ['Case Studies'],
#     'summary': 'Save final summary',
#     'description': 'Save the final case study summary with updated names',
#     'requestBody': {
#         'required': True,
#         'content': {
#             'application/json': {
#                 'schema': {
#                     'type': 'object',
#                     'required': ['case_study_id', 'final_summary'],
#                     'properties': {
#                         'case_study_id': {
#                             'type': 'integer',
#                         'description': 'ID of the case study'
#                         },
#                         'final_summary': {
#                             'type': 'string',
#                             'description': 'Final case study summary text'
#                         },
#                         'solution_provider': {
#                             'type': 'string',
#                             'description': 'Name of the solution provider'
#                         },
#                         'client_name': {
#                             'type': 'string',
#                             'description': 'Name of the client'
#                         },
#                         'project_name': {
#                             'type': 'string',
#                             'description': 'Name of the project'
#                         }
#                     }
#                 }
#             }
#         }
#     },
#     'responses': {
#         200: {
#             'description': 'Final summary saved successfully',
#             'content': {
#                 'application/json': {
#                     'schema': {
#                         'type': 'object',
#                         'properties': {
#                             'status': {'type': 'string'},
#                             'message': {'type': 'string'},
#                             'names': {'type': 'object'}
#                         }
#                     }
#                 }
#             }
#         },
#         400: {'description': 'Missing required data'},
#         404: {'description': 'Case study not found'}
#     }
# })
# def save_final_summary():
#     """Save final summary"""
#     try:
#         data = request.get_json()
#         case_study_id = data.get("case_study_id")
#         final_summary = data.get("final_summary")
#         # Get edited names from frontend if provided
#         solution_provider = data.get("solution_provider")
#         client_name = data.get("client_name") 
#         project_name = data.get("project_name")
# 
#         if not case_study_id or not final_summary:
#             return jsonify({"status": "error", "message": "Missing data"}), 400
# 
#         # Get the case study from DB
#         case_study = CaseStudy.query.filter_by(id=case_study_id).first()
#         try:
#             return jsonify({"status": "error", "message": "Case study not found"}), 404
# 
#         # Update final summary
#         case_study.final_summary = final_summary
# 
#         # Use edited names from frontend if provided, otherwise extract from final summary
#         if solution_provider and client_name and project_name:
#             names = {
#                 "lead_entity": solution_provider,
#                 "partner_entity": client_name,
#                 "project_title": project_name
#             }
#             print(f"✅ Using edited names from frontend for final summary: {names}")
#         else:
#             # Extract names from the new final summary
#             ai_service = AIService()
#             names = ai_service.extract_names_from_case_study(final_summary)
#             print(f"✅ Using extracted names from final summary: {names}")
#         
#         lead_entity = names["lead_entity"]
#         partner_entity = names["partner_entity"]
#         project_title = names["project_title"]
# 
#         # Update CaseStudy name fields (but DON'T override the title - keep the short hook)
#         # The title was already set correctly in the initial generation with a short hook
#         case_study.provider_name = lead_entity
#         case_study.client_name = partner_entity
#         case_study.project_name = project_title
# 
#         db.session.commit()
# 
#         return jsonify({
#             "status": "success",
#             "message": "Final summary and title updated",
#             "names": names,
#             "case_study_id": case_study.id
#         })
# 
#     except Exception as e:
#         db.session.rollback()
#         return jsonify({"status": "error", "message": str(e)}), 500

# Helper functions
def store_solution_provider_session(provider_session_id, cleaned_case_study):
    """Store solution provider session"""
    try:
        ai_service = AIService()
        extracted_names = ai_service.extract_names_from_case_study(cleaned_case_study)
        user_id = get_current_user_id()
        if not user_id:
            raise Exception('No user is logged in.')
        
        from app.models import User
        user = User.query.filter_by(id=user_id).first()
        if not user:
            raise Exception('Logged-in user not found.')

        # Create the CaseStudy (links to user)
        # Extract the title from the first line of the case study content
        lines = cleaned_case_study.split('\n')
        title = lines[0].strip() if lines else "Case Study"
        
        print(f"🔍 DEBUG: Extracted title from first line: '{title}'")
        print(f"🔍 DEBUG: First few lines of case study:")
        for i, line in enumerate(lines[:3]):
            print(f"   Line {i}: '{line}'")
        
        case_study = CaseStudy(
            user_id=user.id,
            title=title,
            final_summary=None  # We fill this later, after full doc is generated
        )
        db.session.add(case_study)
        db.session.commit()

        # Create the SolutionProviderInterview
        provider_interview = SolutionProviderInterview(
            case_study_id=case_study.id,
            session_id=provider_session_id,
            transcript="",  # You can store transcript here later if needed
            summary=cleaned_case_study
        )
        db.session.add(provider_interview)
        db.session.commit()

        return case_study.id  # Return case_study.id to be used for next step

    except Exception as e:
        db.session.rollback()
        raise

def create_client_session(case_study_id):
    """Create client session"""
    try:
        token = str(uuid.uuid4())
        invite_token = InviteToken(
            case_study_id=case_study_id,
            token=token,
            used=False
        )
        db.session.add(invite_token)
        db.session.commit()
        print(f"✅ Client invite token created: {token}")
        return token
    except Exception as e:
        db.session.rollback()
        print("❌ Error creating client invite token:", str(e))
        return None 

@bp.route('/health')
@swag_from({
    'tags': ['System'],
    'summary': 'Health check',
    'description': 'General health check endpoint for debugging',
    'responses': {
        200: {
            'description': 'Health status retrieved successfully',
            'content': {
                'application/json': {
                    'schema': {
                        'type': 'object',
                        'properties': {
                            'status': {'type': 'string'},
                            'database': {'type': 'string'},
                            'environment': {
                                'type': 'object',
                                'properties': {
                                    'flask_env': {'type': 'string'},
                                    'database_url_set': {'type': 'boolean'},
                                    'secret_key_set': {'type': 'boolean'}
                                }
                            }
                        }
                    }
                }
            }
        }
    }
})
def health_check():
    """Health check endpoint for debugging"""
    try:
        # Test database connection
        db.engine.execute("SELECT 1")
        db_status = "healthy"
    except Exception as e:
        logger.error(f"Database health check failed: {str(e)}")
        db_status = f"unhealthy: {str(e)}"
    
    return jsonify({
        "status": "ok",
        "database": db_status,
        "environment": {
            "flask_env": request.environ.get('FLASK_ENV', 'unknown'),
            "database_url_set": bool(request.environ.get('DATABASE_URL')),
            "secret_key_set": bool(request.environ.get('SECRET_KEY'))
        }
    })

@bp.route('/api/health')
@swag_from({
    'tags': ['System'],
    'summary': 'API health check',
    'description': 'API health check endpoint',
    'responses': {
        200: {
            'description': 'API health status',
            'content': {
                'application/json': {
                    'schema': {
                        'type': 'object',
                        'properties': {
                            'status': {'type': 'string'},
                            'message': {'type': 'string'}
                        }
                    }
                }
            }
        }
    }
})
def api_health_check():
    """API health check endpoint"""
    return jsonify({"status": "healthy", "message": "API is running"})

@bp.route('/api/db-health')
@swag_from({
    'tags': ['System'],
    'summary': 'Database health check',
    'description': 'Database health check endpoint',
    'responses': {
        200: {
            'description': 'Database health status',
            'content': {
                'application/json': {
                    'schema': {
                        'type': 'object',
                        'properties': {
                            'status': {'type': 'string'},
                            'database': {'type': 'string'},
                            'users_table_exists': {'type': 'boolean'},
                            'available_tables': {
                                'type': 'array',
                                'items': {'type': 'string'}
                            }
                        }
                    }
                }
            }
        },
        500: {
            'description': 'Database health check failed',
            'content': {
                'application/json': {
                    'schema': {
                        'type': 'object',
                        'properties': {
                            'status': {'type': 'string'},
                            'database': {'type': 'string'},
                            'error': {'type': 'string'},
                            'error_type': {'type': 'string'}
                        }
                    }
                }
            }
        }
    }
})
def db_health_check():
    """Database health check endpoint"""
    try:
        # Test database connection
        db.engine.execute("SELECT 1")
        
        # Check if users table exists
        tables = db.engine.table_names()
        users_table_exists = 'users' in tables
        
        return jsonify({
            "status": "healthy",
            "database": "connected",
            "users_table_exists": users_table_exists,
            "available_tables": tables
        })
    except Exception as e:
        return jsonify({
            "status": "unhealthy",
            "database": "error",
            "error": str(e),
            "error_type": str(type(e))
        }), 500

@bp.route("/download/<filename>")
@swag_from({
    'tags': ['Files'],
    'summary': 'Download file',
    'description': 'Download generated PDF or Word file',
    'parameters': [
        {
            'name': 'filename',
            'in': 'path',
            'required': True,
            'schema': {'type': 'string'},
            'description': 'Name of the file to download'
        }
    ],
    'responses': {
        200: {
            'description': 'File downloaded successfully',
            'content': {
                'application/octet-stream': {
                    'schema': {
                        'type': 'string',
                        'format': 'binary'
                    }
                }
            }
        },
        404: {'description': 'File not found'}
    }
})
def download_file(filename):
    """Download generated PDF or Word file"""
    try:
        import os
        file_path = f"generated_pdfs/{filename}"
        print(f"🔍 DEBUG: Attempting to download file: {file_path}")
        print(f"🔍 DEBUG: File exists: {os.path.exists(file_path)}")
        
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found"}), 404
            
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        print(f"🔍 ERROR: Download failed: {str(e)}")
        return jsonify({"error": str(e)}), 500

@bp.route("/download_word/<int:case_study_id>")
@swag_from({
    'tags': ['Files'],
    'summary': 'Download Word document',
    'description': 'Download Word document for a specific case study',
    'parameters': [
        {
            'name': 'case_study_id',
            'in': 'path',
            'required': True,
            'schema': {'type': 'integer'},
            'description': 'ID of the case study'
        }
    ],
    'responses': {
        200: {
            'description': 'Word document downloaded successfully',
            'content': {
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': {
                    'schema': {
                        'type': 'string',
                        'format': 'binary'
                    }
                }
            }
        },
        400: {'description': 'No final summary available'},
        404: {'description': 'Case study not found'}
    }
})
def download_word_document(case_study_id):
    """Download Word document for a specific case study"""
    try:
        print("WORD called")
        from app.models import CaseStudy
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        
        case_study = CaseStudy.query.filter_by(id=case_study_id).first()
        if not case_study:
            return jsonify({"error": "Case study not found"}), 404
            
        if not case_study.final_summary:
            return jsonify({"error": "No final summary available"}), 400

        # Create Word document
        doc = Document()
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(case_study.title or "Case Study")
        title_run.bold = True
        title_run.font.size = Pt(20)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some spacing
        doc.add_paragraph()
        
        # Add the final summary content
        lines = case_study.final_summary.split('\n')
        for line in lines:
            if line.strip():  # Only add non-empty lines
                # Check if it's a header (all caps or starts with **)
                if line.strip().isupper() or line.strip().startswith('**'):
                    # It's a header
                    header_para = doc.add_paragraph()
                    header_run = header_para.add_run(line.strip().replace('**', ''))
                    header_run.bold = True
                    header_run.font.size = Pt(15)
                else:
                    # It's regular content
                    para = doc.add_paragraph()
                    para.add_run(line.strip())
        
        # Save to BytesIO buffer
        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        
        return send_file(
            word_buffer,
            as_attachment=True,
            download_name=f"{case_study.title or 'Case_Study'}_{case_study_id}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"🔍 ERROR: Word document generation failed: {str(e)}")
        return jsonify({"error": str(e)}), 500 